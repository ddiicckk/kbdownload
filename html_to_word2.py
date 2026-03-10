import os
import re
import pandas as pd
from docx import Document
from html2docx import html2docx

# ==========================
# Configuration
# ==========================
EXCEL_PATH = "kb_knowledge.xlsx"     # <-- your Excel path
SHEET_NAME = 0                       # index or sheet name
COL_ID = "Number"                    # filename source
COL_HTML = "Article body"            # HTML source
OUTPUT_DIR = "exported_docs"         # output folder
LOG_PATH = "export_log.txt"          # optional text log

# ==========================
# Helpers
# ==========================
def sanitize_filename(name: str) -> str:
    name = str(name).strip()
    name = re.sub(r'[<>:"/\\|?*]', "_", name)
    name = re.sub(r"\s+", " ", name)
    return name[:200] if len(name) > 200 else name

# Remove XML-illegal control chars that may appear in pasted HTML
INVALID_XML_RE = re.compile(
    r"(?:"                       # any of:
    r"[\x00-\x08]"               # C0 controls
    r"|[\x0B-\x0C]"
    r"|[\x0E-\x1F]"
    r"|[\uD800-\uDFFF]"          # surrogate blocks
    r"|[\uFFFE\uFFFF]"           # non-characters
    r")"
)

def clean_html(html: str) -> str:
    if html is None:
        return ""
    s = str(html).replace("\r\n", "\n").replace("\r", "\n")
    s = INVALID_XML_RE.sub("", s)
    # Wrap fragments in a minimal document so parsers behave better
    if "<html" not in s.lower():
        s = (
            "<!DOCTYPE html>"
            "<html><head><meta charset='utf-8'></head><body>"
            f"{s}"
            "</body></html>"
        )
    return s

def strip_html_to_text(s: str) -> str:
    if not isinstance(s, str):
        s = str(s) if s is not None else ""
    s = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", s)  # remove script/style
    s = re.sub(r"(?s)<[^>]+>", "", s)                      # strip tags
    s = s.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    s = re.sub(r"\s+\n", "\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()

# ==========================
# Main
# ==========================
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    with open(LOG_PATH, "w", encoding="utf-8") as logf:
        logf.write("Export run log\n")

    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, engine="openpyxl")

    missing = [c for c in (COL_ID, COL_HTML) if c not in df.columns]
    if missing:
        raise ValueError(f"Missing required column(s): {missing}")

    exported_ok = 0
    exported_fallback = 0
    skipped = 0
    failures = 0

    for idx, row in df.iterrows():
        excel_row_number = idx + 2  # header row is 1; pandas index 0 -> Excel row 2

        file_id = row.get(COL_ID)
        html_raw = row.get(COL_HTML)

        if pd.isna(file_id) or pd.isna(html_raw) or str(html_raw).strip() == "":
            skipped += 1
            with open(LOG_PATH, "a", encoding="utf-8") as logf:
                logf.write(f"[SKIP] Excel row {excel_row_number}: missing Number or Article body\n")
            continue

        safe_name = sanitize_filename(file_id)
        out_path = os.path.join(OUTPUT_DIR, f"{safe_name}.docx")

        try:
            # Clean + convert HTML → docx
            cleaned = clean_html(str(html_raw))
            doc = Document()
            html2docx(cleaned, doc)
            doc.save(out_path)
            exported_ok += 1
            print(f"[OK]    {excel_row_number} -> {out_path}")
            with open(LOG_PATH, "a", encoding="utf-8") as logf:
                logf.write(f"[OK] Excel row {excel_row_number} '{safe_name}.docx'\n")
        except Exception as e:
            # Fallback: store as plain text so nothing is lost
            try:
                text = strip_html_to_text(str(html_raw)) or "(No renderable text after HTML cleanup)"
                doc = Document()
                doc.add_paragraph(f"[Fallback: plain text render for '{safe_name}']")
                doc.add_paragraph("")
                for para in text.splitlines():
                    doc.add_paragraph(para)
                doc.save(out_path)
                exported_fallback += 1
                print(f"[TXT]   {excel_row_number} -> {out_path} (fallback)")
                with open(LOG_PATH, "a", encoding="utf-8") as logf:
                    logf.write(f"[TXT] Excel row {excel_row_number} '{safe_name}.docx' (fallback) due to: {e}\n")
            except Exception as e2:
                failures += 1
                print(f"[FAIL]  {excel_row_number} -> {safe_name}.docx | {e2}")
                with open(LOG_PATH, "a", encoding="utf-8") as logf:
                    logf.write(f"[FAIL] Excel row {excel_row_number} id '{safe_name}': {e} | fallback error: {e2}\n")

    print("\n=== Summary ===")
    print(f"Converted (HTML): {exported_ok}")
    print(f"Fallback (Text) : {exported_fallback}")
    print(f"Skipped (Empty) : {skipped}")
    print(f"Failures        : {failures}")
    print(f"Log             : {LOG_PATH}")

if __name__ == "__main__":
    main()
``
