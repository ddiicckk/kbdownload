# -*- coding: utf-8 -*-
"""
Export each Excel row to a Word (.docx) file:
- File name  : value in column 'Number'
- File content: value in column 'Article body' (HTML -> native Word content)

Requires:
    pip install pandas openpyxl python-docx beautifulsoup4 lxml
"""

import os
import re
import uuid
from typing import Optional, Union

import pandas as pd
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================
# Configuration
# ==========================
EXCEL_PATH = "kb_knowledge.xlsx"  # <-- your Excel file
SHEET_NAME = 0                    # index (0-based) or sheet name
COL_ID = "Number"                 # filename source
COL_HTML = "Article body"         # HTML source
OUTPUT_DIR = "exported_docs"
LOG_PATH = "export_log.txt"

# If your HTML is very large and heavily nested tables, you can cap column widths or truncate text if needed.
MAX_PARAGRAPH_LENGTH = None  # set to e.g. 20000 to cap any single paragraph text (None = no cap)

# ==========================
# Utilities
# ==========================
INVALID_XML_RE = re.compile(
    r"(?:[\x00-\x08]|[\x0B-\x0C]|[\x0E-\x1F]|[\uD800-\uDFFF]|[\uFFFE\uFFFF])"
)

def clean_text_for_xml(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = INVALID_XML_RE.sub("", s)
    # Fix a few common entities
    s = s.replace("\xa0", " ").replace("&nbsp;", " ")
    return s

def sanitize_filename(name: Union[str, int]) -> str:
    name = clean_text_for_xml(str(name)).strip()
    name = re.sub(r'[<>:"/\\|?*]', "_", name)
    name = re.sub(r"\s+", " ", name)
    return name[:200]

def ensure_unique_path(base_dir: str, base_name: str, ext: str = ".docx") -> str:
    """
    Returns a non-clashing path by appending (2), (3), ...
    """
    candidate = os.path.join(base_dir, f"{base_name}{ext}")
    if not os.path.exists(candidate):
        return candidate
    i = 2
    while True:
        candidate = os.path.join(base_dir, f"{base_name} ({i}){ext}")
        if not os.path.exists(candidate):
            return candidate
        i += 1

def add_hyperlink(paragraph, url: str, text: str, underline=True, color="0000FF"):
    """
    Create a hyperlink within a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Build the run
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    # Text node
    t = OxmlElement("w:t")
    t.text = text

    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)

    # Append to paragraph
    paragraph._p.append(hyperlink)

# ==========================
# HTML -> Word conversion
# ==========================
class HtmlToDocx:
    """
    Minimal-yet-robust HTML -> python-docx converter for common tags:
    headings, paragraphs/divs, b/i/u, br, lists, tables, and links.
    """

    def __init__(self, doc: Document):
        self.doc = doc
        self.list_stack = []  # track nested list types: "ul"/"ol"

    def convert(self, html: str):
        if html is None:
            return
        s = clean_text_for_xml(str(html))
        # Some sources escape < and > as \"\<p\>\"; try to un-escape if we see that pattern
        s = s.replace("\\<", "<").replace("\\>", ">")

        # Wrap fragment if needed
        if "<html" not in s.lower():
            s = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{s}</body></html>"

        # Parse
        # Prefer lxml if installed; BeautifulSoup will pick it, else fallback to 'html.parser'
        soup = BeautifulSoup(s, "lxml")
        body = soup.body or soup

        # Convert children
        for child in body.children:
            self._handle_block(child)

    # --------- Block-level handlers ---------
    def _handle_block(self, node):
        if isinstance(node, NavigableString):
            text = clean_text_for_xml(str(node)).strip()
            if text:
                p = self.doc.add_paragraph()
                self._append_inline(p, node)
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        if name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(name[1])
            text = self._gather_text(node)
            p = self.doc.add_heading(level=level)
            if text:
                # Use runs so inline bold/italics are reflected (heading text property is plain)
                self._append_inline(p, node)
            return

        if name in ["p", "div"]:
            p = self.doc.add_paragraph()
            self._append_inline(p, node)
            return

        if name in ["ul", "ol"]:
            self.list_stack.append(name)  # push
            for li in node.find_all("li", recursive=False):
                style = "List Bullet" if name == "ul" else "List Number"
                p = self.doc.add_paragraph(style=style)
                self._append_inline(p, li)
            self.list_stack.pop()  # pop
            return

        if name == "table":
            self._handle_table(node)
            return

        if name in ["br", "hr"]:
            # force a paragraph break for hr, line break handled at inline level for br
            if name == "hr":
                self.doc.add_paragraph().add_run("_" * 40)
            else:
                self.doc.add_paragraph()
            return

        # Unknown/other block: process children in flow
        for child in node.children:
            self._handle_block(child)

    def _handle_table(self, table_tag: Tag):
        rows = table_tag.find_all("tr", recursive=True)
        if not rows:
            return
        # Determine max columns by scanning first few rows
        max_cols = 0
        grid = []
        for r in rows:
            cells = [c for c in r.find_all(["td", "th"], recursive=False)]
            max_cols = max(max_cols, len(cells))
            grid.append(cells)

        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(grid), cols=max_cols)
        table.style = "Table Grid"

        for r_idx, cells in enumerate(grid):
            for c_idx in range(max_cols):
                cell = table.cell(r_idx, c_idx)
                cell_para = cell.paragraphs[0]
                if c_idx < len(cells):
                    self._append_inline(cell_para, cells[c_idx])
                else:
                    # empty filler
                    cell_para.add_run("")

    # --------- Inline handlers ---------
    def _append_inline(self, paragraph, node, inherited=None):
        """
        Append node's inline contents to an existing paragraph.
        """
        if inherited is None:
            inherited = {"bold": False, "italic": False, "underline": False}

        if isinstance(node, NavigableString):
            txt = clean_text_for_xml(str(node))
            if MAX_PARAGRAPH_LENGTH and len(txt) > MAX_PARAGRAPH_LENGTH:
                txt = txt[:MAX_PARAGRAPH_LENGTH] + "…"
            if txt:
                run = paragraph.add_run(txt)
                run.bold = inherited["bold"]
                run.italic = inherited["italic"]
                run.underline = inherited["underline"]
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        # Compute formatting inheritance
        fmt = dict(inherited)
        if name in ["strong", "b"]:
            fmt["bold"] = True
        if name in ["em", "i"]:
            fmt["italic"] = True
        if name in ["u"]:
            fmt["underline"] = True

        # Also honor <span style="font-weight:bold">, etc.
        style_attr = node.get("style", "")
        if style_attr:
            if re.search(r"font-weight\s*:\s*(bold|700|800|900)", style_attr, re.I):
                fmt["bold"] = True
            if re.search(r"font-style\s*:\s*italic", style_attr, re.I):
                fmt["italic"] = True
            if re.search(r"text-decoration\s*:\s*underline", style_attr, re.I):
                fmt["underline"] = True

        if name == "br":
            paragraph.add_run().add_break()
            return

        if name == "a":
            href = node.get("href", "").strip()
            link_text = self._gather_text(node) or href
            if href:
                add_hyperlink(paragraph, href, link_text)
            else:
                paragraph.add_run(link_text)
            return

        if name in ["img"]:
            # Many source images are relative (/sys_attachment...) and not fetchable here.
            # We drop in a placeholder with the source URL so you can post-process if needed.
            src = node.get("src")
            alt = node.get("alt", "") or ""
            placeholder = f"[Image: {alt}] ({src})" if src else f"[Image: {alt}]"
            paragraph.add_run(placeholder)
            return

        # Default: walk children
        for child in node.children:
            self._append_inline(paragraph, child, fmt)

    def _gather_text(self, node) -> str:
        return clean_text_for_xml("".join(node.stripped_strings))


# ==========================
# Main
# ==========================
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    with open(LOG_PATH, "w", encoding="utf-8") as logf:
        logf.write("Export run log\n")

    # Read Excel
    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, engine="openpyxl")

    # Validate columns
    missing = [c for c in (COL_ID, COL_HTML) if c not in df.columns]
    if missing:
        raise ValueError(f"Missing required column(s): {missing}")

    exported = 0
    skipped = 0
    failures = 0

    for idx, row in df.iterrows():
        excel_row_number = idx + 2  # header is row 1
        number = row.get(COL_ID)
        html = row.get(COL_HTML)

        if pd.isna(number) or pd.isna(html) or str(html).strip() == "":
            skipped += 1
            with open(LOG_PATH, "a", encoding="utf-8") as logf:
                logf.write(f"[SKIP] Excel row {excel_row_number}: missing Number or Article body\n")
            continue

        base_name = sanitize_filename(number)
        out_path = ensure_unique_path(OUTPUT_DIR, base_name, ext=".docx")

        try:
            doc = Document()

            # (Optional) Title: uncomment if you want the Number at top of the doc
            # doc.add_heading(str(number), level=1)

            # Set base font (optional)
            for style_name in ["Normal", "List Paragraph"]:
                try:
                    doc.styles[style_name].font.name = "Calibri"
                    doc.styles[style_name].font.size = Pt(11)
                except Exception:
                    pass

            converter = HtmlToDocx(doc)
            converter.convert(str(html))

            # Ensure doc isn't empty (in extreme edge cases)
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                doc.add_paragraph("(No renderable content)")

            doc.save(out_path)
            exported += 1

            with open(LOG_PATH, "a", encoding="utf-8") as logf:
                logf.write(f"[OK] Excel row {excel_row_number} -> '{os.path.basename(out_path)}'\n")

        except Exception as e:
            failures += 1
            with open(LOG_PATH, "a", encoding="utf-8") as logf:
                logf.write(f"[FAIL] Excel row {excel_row_number} id '{number}': {e}\n")

    # Summary
    print("\n=== Summary ===")
    print(f"Exported      : {exported}")
    print(f"Skipped (empty): {skipped}")
    print(f"Failures      : {failures}")
    print(f"Output folder : {OUTPUT_DIR}")
    print(f"Log           : {LOG_PATH}")

if __name__ == "__main__":
    main()
