# -*- coding: utf-8 -*-
"""
Convert CSV rows into Word (.docx) files.
- 'number' column --> output filename
- 'text'   column --> HTML content

This script:
  ✔ Handles unlimited-length HTML (CSV has no 32K limit)
  ✔ Converts HTML to real Word formatting (not altChunk)
  ✔ Supports headings, paragraphs, lists, tables, bold/italic/underline, links
  ✔ Avoids overwritten files by auto-suffixing duplicates
  ✔ Produces a log
"""

import os
import re
import pandas as pd
from typing import Union
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================
# CONFIGURATION
# ==========================
CSV_PATH = "kb_articles.csv"          # <-- update if needed
COL_ID = "number"                     # filename source
COL_HTML = "text"                     # HTML source
OUTPUT_DIR = "exported_docs"
LOG_PATH = "export_log.txt"

# ==========================
# UTILITIES
# ==========================
INVALID_XML_RE = re.compile(
    r"(?:[\x00-\x08]|[\x0B-\x0C]|[\x0E-\x1F]|[\uD800-\uDFFF]|[\uFFFE\uFFFF])"
)

def clean_text_for_xml(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = INVALID_XML_RE.sub("", s)
    s = s.replace("\xa0", " ").replace("&nbsp;", " ")
    return s

def sanitize_filename(name: Union[str, int]) -> str:
    name = clean_text_for_xml(str(name)).strip()
    name = re.sub(r'[<>:"/\\|?*]', "_", name)
    name = re.sub(r"\s+", " ", name)
    return name[:200]

def ensure_unique_path(base_dir: str, base_name: str, ext: str = ".docx") -> str:
    """Generate a path without overwriting existing files."""
    candidate = os.path.join(base_dir, f"{base_name}{ext}")
    if not os.path.exists(candidate):
        return candidate
    i = 2
    while True:
        candidate = os.path.join(base_dir, f"{base_name} ({i}){ext}")
        if not os.path.exists(candidate):
            return candidate
        i += 1

def add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink into a Word paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)

    t = OxmlElement("w:t")
    t.text = text

    run.append(rPr)
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

# ==========================
# HTML → Word converter
# ==========================
class HtmlToDocx:
    def __init__(self, doc: Document):
        self.doc = doc

    def convert(self, html: str):
        html = clean_text_for_xml(html)

        # wrap if fragment
        if "<html" not in html.lower():
            html = f"<html><body>{html}</body></html>"

        soup = BeautifulSoup(html, "lxml")
        body = soup.body or soup

        for child in body.children:
            self._handle_block(child)

    def _handle_block(self, node):
        if isinstance(node, NavigableString):
            text = node.strip()
            if text:
                p = self.doc.add_paragraph(text)
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        if name in ["h1","h2","h3","h4","h5","h6"]:
            level = int(name[1])
            p = self.doc.add_heading(level=level)
            self._append_inline(p, node)
            return

        if name in ["p","div"]:
            p = self.doc.add_paragraph()
            self._append_inline(p, node)
            return

        if name in ["ul","ol"]:
            list_style = "List Bullet" if name=="ul" else "List Number"
            for li in node.find_all("li", recursive=False):
                p = self.doc.add_paragraph(style=list_style)
                self._append_inline(p, li)
            return

        if name == "table":
            self._handle_table(node)
            return

        if name == "br":
            self.doc.add_paragraph()
            return

        for child in node.children:
            self._handle_block(child)

    def _handle_table(self, table_tag: Tag):
        rows = table_tag.find_all("tr", recursive=True)
        if not rows:
            return
        grid = []
        max_cols = 0
        for r in rows:
            cells = r.find_all(["td","th"], recursive=False)
            max_cols = max(max_cols, len(cells))
            grid.append(cells)

        table = self.doc.add_table(rows=len(grid), cols=max_cols)
        table.style = "Table Grid"

        for r_idx, row_cells in enumerate(grid):
            for c_idx in range(max_cols):
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                if c_idx < len(row_cells):
                    self._append_inline(p, row_cells[c_idx])
                else:
                    p.add_run("")

    def _append_inline(self, paragraph, node, fmt=None):
        if fmt is None:
            fmt = {"bold": False, "italic": False, "underline": False}

        if isinstance(node, NavigableString):
            text = clean_text_for_xml(str(node))
            r = paragraph.add_run(text)
            r.bold = fmt["bold"]
            r.italic = fmt["italic"]
            r.underline = fmt["underline"]
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        local_fmt = dict(fmt)

        if name in ["strong","b"]:
            local_fmt["bold"] = True
        if name in ["em","i"]:
            local_fmt["italic"] = True
        if name == "u":
            local_fmt["underline"] = True

        style_attr = node.get("style","")
        if "font-weight" in style_attr and "bold" in style_attr:
            local_fmt["bold"] = True
        if "font-style" in style_attr and "italic" in style_attr:
            local_fmt["italic"] = True
        if "text-decoration" in style_attr and "underline" in style_attr:
            local_fmt["underline"] = True

        if name == "br":
            paragraph.add_run().add_break()
            return

        if name == "a":
            href = node.get("href","").strip()
            text = node.get_text(strip=True) or href
            if href:
                add_hyperlink(paragraph, href, text)
            else:
                paragraph.add_run(text)
            return

        if name == "img":
            src = node.get("src","")
            alt = node.get("alt","")
            paragraph.add_run(f"[Image: {alt}] ({src})")
            return

        for child in node.children:
            self._append_inline(paragraph, child, local_fmt)

# ==========================
# MAIN
# ==========================
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    with open(LOG_PATH,"w",encoding="utf-8") as log:
        log.write("CSV → DOCX Export Log\n")

    df = pd.read_csv(CSV_PATH)

    missing = [c for c in (COL_ID, COL_HTML) if c not in df.columns]
    if missing:
        raise ValueError(f"Missing required columns: {missing}")

    exported = 0
    skipped = 0
    failures = 0

    for idx, row in df.iterrows():
        file_id = row.get(COL_ID)
        html = row.get(COL_HTML)

        if pd.isna(file_id) or pd.isna(html) or str(html).strip()=="":
            skipped += 1
            with open(LOG_PATH,"a",encoding="utf-8") as log:
                log.write(f"[SKIP] Row {idx+1}: empty id/text\n")
            continue

        base_name = sanitize_filename(file_id)
        out_path = ensure_unique_path(OUTPUT_DIR, base_name, ".docx")

        try:
            doc = Document()
            converter = HtmlToDocx(doc)
            converter.convert(str(html))

            if len(doc.paragraphs)==0 and len(doc.tables)==0:
                doc.add_paragraph("(No renderable content)")

            doc.save(out_path)
            exported += 1

            with open(LOG_PATH,"a",encoding="utf-8") as log:
                log.write(f"[OK] {file_id} -> {out_path}\n")

        except Exception as e:
            failures += 1
            with open(LOG_PATH,"a",encoding="utf-8") as log:
                log.write(f"[FAIL] Row {idx+1} id '{file_id}': {e}\n")

    print("=== SUMMARY ===")
    print("Exported:", exported)
    print("Skipped:", skipped)
    print("Failed :", failures)
    print("Output :", OUTPUT_DIR)
    print("Log    :", LOG_PATH)

if __name__ == "__main__":
    main()
